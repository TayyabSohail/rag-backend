import fitz
import docx
import html2text
import requests

def parse_and_chunk(content: bytes, filename: str):
    chunks = []

    if filename.endswith(".pdf"):
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc:
            text = page.get_text()
            chunks.extend(chunk_text(text))

    elif filename.endswith(".docx"):
        doc = docx.Document()
        with open("temp.docx", "wb") as f:
            f.write(content)
        doc = docx.Document("temp.docx")
        text = "\n".join([para.text for para in doc.paragraphs])
        chunks.extend(chunk_text(text))

    elif filename.endswith(".html") or filename.endswith(".htm"):
        text = html2text.html2text(content.decode("utf-8"))
        chunks.extend(chunk_text(text))

    return chunks

def chunk_text(text: str, max_tokens: int = 300):
    sentences = text.split(". ")
    chunks, chunk = [], ""
    for sentence in sentences:
        if len(chunk) + len(sentence) < max_tokens:
            chunk += sentence + ". "
        else:
            chunks.append(chunk.strip())
            chunk = sentence + ". "
    if chunk:
        chunks.append(chunk.strip())
    return chunks

def parse_and_chunk_from_url(url: str):
    try:
        response = requests.get(url, timeout=10)
        if response.status_code != 200:
            raise Exception("Failed to fetch URL")
        html_content = response.text
        text = html2text.html2text(html_content)
        return chunk_text(text)
    except Exception as e:
        return [f"Error fetching or parsing URL: {str(e)}"]
